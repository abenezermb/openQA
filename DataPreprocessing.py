# input
# .pdf,.docx,.txt

# Operation
# read the text from the pdf, docx and txt files

# Output
# returns a list of strings(every paragraph from every document is converted to an element of the list)

def readPdf(filename):
    import pdfplumber
    pages_text = []
    with pdfplumber.open(filename) as pdf:
        for i in range(len(pdf.pages)):
            page = pdf.pages[i]
            # print(page.extract_text())
            pages_text.append(page.extract_text())
        return pages_text
def readDocx(filename):
    # import docx
    from docx import Document
    try:
        doc = Document(filename) # word reader object
        docs = []
        for i in range(len(doc.paragraphs)):
            docs.append(doc.paragraphs[i].text) 
        return docs
    except IOError:
        print("There was an error opening this word file")
        return

# # read a text from a word  file
# filename = "./sample.docx"
# text = readDocx(filename)
# print(text)
# read the text from a .pdf file

# filename = './sample.pdf'
# text = fetchText(filename)
# print(text)